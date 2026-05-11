"""
app/parsers/docx_parser.py
===========================
DOCX text extraction using python-docx.
Handles paragraphs, tables, and headers.
"""

from pathlib import Path
from typing import List
from app.utils.logger import get_logger

logger = get_logger(__name__)


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract all text content from a DOCX file.
    Includes paragraphs, table cells, and styled content.

    Args:
        file_path: Absolute path to the .docx file.

    Returns:
        Extracted text as a single string.

    Raises:
        ValueError: If file doesn't exist.
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    path = Path(file_path)
    if not path.exists():
        raise ValueError(f"DOCX file not found: {file_path}")

    doc = Document(file_path)
    text_parts: List[str] = []

    # --- Extract paragraphs ---
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())

    # --- Extract table content ---
    for table in doc.tables:
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_texts.append(cell_text)
            if row_texts:
                text_parts.append(" | ".join(row_texts))

    extracted = "\n".join(text_parts)
    logger.info(
        f"DOCX extracted {len(extracted)} chars from {path.name}"
    )
    return extracted


def extract_text_from_docx_bytes(file_bytes: bytes, filename: str = "upload.docx") -> str:
    """
    Extract text from DOCX bytes (for in-memory uploads).

    Args:
        file_bytes: Raw DOCX bytes.
        filename: Hint for logging.

    Returns:
        Extracted text string.
    """
    import io

    try:
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        text_parts: List[str] = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())

        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_texts:
                    text_parts.append(" | ".join(row_texts))

        extracted = "\n".join(text_parts)
        logger.debug(f"DOCX bytes extracted {len(extracted)} chars ({filename})")
        return extracted

    except Exception as e:
        logger.error(f"DOCX bytes extraction failed for {filename}: {e}")
        return ""
